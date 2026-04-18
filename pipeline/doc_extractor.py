"""
문서 파일(Word/PDF/Excel/XML)에서 텍스트 라인 추출
- .docx : python-docx (단락 + 표 셀)
- .pdf  : pdfplumber (텍스트 레이아웃 유지)
- .xlsx/.xls : openpyxl (각 행 → 탭 구분 텍스트)
- .xml  : ElementTree (요소 텍스트 + 속성값 → 구조화 라인)
"""
from __future__ import annotations
from pathlib import Path
from typing import Iterator

SUPPORTED_DOC_EXTENSIONS = {'.docx', '.pdf', '.xlsx', '.xls', '.xlsm', '.xml'}


def is_document(file_path: str | Path) -> bool:
    return Path(file_path).suffix.lower() in SUPPORTED_DOC_EXTENSIONS


def extract_lines(file_path: str | Path) -> Iterator[tuple[int, str]]:
    """
    문서에서 텍스트 라인을 추출합니다.
    Yields: (line_no, text)
    """
    path = Path(file_path)
    ext  = path.suffix.lower()

    if ext == '.docx':
        yield from _extract_docx(path)
    elif ext == '.pdf':
        yield from _extract_pdf(path)
    elif ext in ('.xlsx', '.xls', '.xlsm'):
        yield from _extract_excel(path)
    elif ext == '.xml':
        yield from _extract_xml(path)


def _extract_docx(path: Path) -> Iterator[tuple[int, str]]:
    try:
        import docx as python_docx
        doc = python_docx.Document(str(path))
        line_no = 1

        # 단락
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                yield line_no, text
                line_no += 1

        # 표 셀
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    yield line_no, '\t'.join(cells)
                    line_no += 1

    except Exception as e:
        print(f"  [DOCX 오류] {path.name}: {e}")


def _extract_pdf(path: Path) -> Iterator[tuple[int, str]]:
    try:
        import pdfplumber
        line_no = 1
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if not text:
                    continue
                for raw_line in text.splitlines():
                    stripped = raw_line.strip()
                    if stripped:
                        yield line_no, stripped
                        line_no += 1
    except Exception as e:
        print(f"  [PDF 오류] {path.name}: {e}")


def _extract_excel(path: Path) -> Iterator[tuple[int, str]]:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        line_no = 1

        for sheet in wb.worksheets:
            # 시트 이름을 헤더로 삽입 (파서가 컨텍스트로 활용 가능)
            yield line_no, f"[Sheet: {sheet.title}]"
            line_no += 1

            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else '' for c in row]
                # 빈 행 스킵
                if not any(c.strip() for c in cells):
                    continue
                yield line_no, '\t'.join(cells)
                line_no += 1

        wb.close()
    except Exception as e:
        print(f"  [Excel 오류] {path.name}: {e}")


def _extract_xml(path: Path) -> Iterator[tuple[int, str]]:
    """
    XML 파일을 파싱하여 PII 탐지용 텍스트 라인을 생성합니다.

    전략:
    1. 리프 요소(자식 없음): 태그명 + 속성 + 텍스트를 한 줄로 합성
       예: <phone>010-1234-5678</phone>  →  "phone: 010-1234-5678"
    2. 속성에 PII가 있을 수 있으므로 속성 key=value도 포함
    3. Log4j XML / 커스텀 감사로그 등 이벤트 블록은 자식 요소 전체를
       한 줄로 합쳐서 로그 라인처럼 재구성
       예: <event timestamp="..." user="emp003"><query>SELECT ...</query></event>
    """
    import xml.etree.ElementTree as ET

    try:
        tree = ET.parse(str(path))
        root = tree.getroot()
    except ET.ParseError:
        # 파싱 실패 시 원본 텍스트 파일처럼 라인 단위로 yield
        try:
            encodings = ['utf-8', 'utf-8-sig', 'euc-kr', 'cp949', 'latin-1']
            for enc in encodings:
                try:
                    text = path.read_text(encoding=enc)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                text = path.read_text(encoding='utf-8', errors='replace')
            for line_no, line in enumerate(text.splitlines(), start=1):
                if line.strip():
                    yield line_no, line.strip()
        except Exception as e2:
            print(f"  [XML 파싱 오류] {path.name}: {e2}")
        return
    except Exception as e:
        print(f"  [XML 오류] {path.name}: {e}")
        return

    line_no = 1

    # 태그명에서 네임스페이스 제거: {http://...}tag → tag
    def local_tag(tag: str) -> str:
        return tag.split('}', 1)[-1] if '}' in tag else tag

    def elem_to_line(elem) -> str:
        """요소 하나를 "tag attr=val ... text" 형태 한 줄로 변환."""
        parts = [local_tag(elem.tag)]
        for k, v in elem.attrib.items():
            parts.append(f'{local_tag(k)}={v}')
        text = (elem.text or '').strip()
        if text:
            parts.append(text)
        return '  '.join(parts)

    def walk(elem, depth: int = 0) -> Iterator[tuple[int, str]]:
        nonlocal line_no
        children = list(elem)

        if not children:
            # 리프 노드: 텍스트 또는 속성이 있는 경우만 출력
            line = elem_to_line(elem)
            if len(line) > len(local_tag(elem.tag)):  # 내용 있을 때만
                yield line_no, line
                line_no += 1
        else:
            # 자식 있는 노드: 속성을 가진 컨테이너면 속성 라인 먼저 출력
            if elem.attrib:
                attr_line = local_tag(elem.tag) + '  ' + '  '.join(
                    f'{local_tag(k)}={v}' for k, v in elem.attrib.items()
                )
                yield line_no, attr_line
                line_no += 1

            # 자식 요소들을 "이벤트 블록"으로 묶어 한 줄 합성 시도
            # (Log4j XML, 감사로그 등에서 유용)
            child_texts = []
            for child in children:
                ct = (child.text or '').strip()
                if ct:
                    child_texts.append(f'{local_tag(child.tag)}={ct}')
                for k, v in child.attrib.items():
                    child_texts.append(f'{local_tag(k)}={v}')

            if child_texts:
                combined = local_tag(elem.tag) + '  ' + '  '.join(child_texts)
                yield line_no, combined
                line_no += 1

            # 자식 각각도 재귀 처리 (중첩 PII 탐지)
            for child in children:
                yield from walk(child, depth + 1)

    yield from walk(root)
